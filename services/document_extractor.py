"""
Document text extraction service for MetricFlow AI.
Supports PDF, DOCX, TXT, and MD file formats.
"""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Service to extract text content from various document formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}

    def extract_text(self, file_content: bytes, filename: str) -> Optional[str]:
        """
        Extract text from a document file.

        Args:
            file_content: The binary content of the file
            filename: The name of the file (used to determine format)

        Returns:
            Extracted text content or None if extraction fails
        """
        extension = self._get_extension(filename)

        if extension not in self.SUPPORTED_EXTENSIONS:
            logger.warning(f"Unsupported file extension: {extension}")
            return None

        try:
            if extension == ".pdf":
                return self._extract_from_pdf(file_content)
            elif extension in (".docx", ".doc"):
                return self._extract_from_docx(file_content)
            elif extension in (".txt", ".md"):
                return self._extract_from_text(file_content)
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            return None

        return None

    def _get_extension(self, filename: str) -> str:
        """Get lowercase file extension."""
        if "." in filename:
            return "." + filename.rsplit(".", 1)[-1].lower()
        return ""

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                logger.error("No PDF library available. Install pypdf or PyPDF2.")
                raise ImportError("PDF extraction requires pypdf or PyPDF2")

        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return "\n\n".join(text_parts)

    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Install it with: pip install python-docx")
            raise ImportError("DOCX extraction requires python-docx")

        docx_file = io.BytesIO(file_content)
        document = Document(docx_file)

        text_parts = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    def _extract_from_text(self, file_content: bytes) -> str:
        """Extract text from TXT or MD file."""
        # Try different encodings
        encodings = ["utf-8", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue

        # Fallback with error handling
        return file_content.decode("utf-8", errors="replace")

    def get_supported_formats(self) -> list[str]:
        """Return list of supported file formats for display."""
        return ["PDF", "Word (DOCX)", "Texto (TXT)", "Markdown (MD)"]
